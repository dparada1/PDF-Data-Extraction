import re
from docx.enum.text import WD_COLOR_INDEX


def convert_height(height):
    """
    Converts height to both feet/inches and centimeters.

    Args:
        height (str): Client's height in various formats.

    Returns:
        str: Height in feet/inches and centimeters.
    """
    height = height.replace(" ", "").strip()
    if height == "M":
        return "M"
    
    height = re.sub(r'[^0-9a-zA-Z]+', '', height)
    
    try:
        if len(height) == 3 and float(height) < 300:
            feet = int(float(height) // 30.48)
            inches = int((float(height) % 30.48) / 2.54)
            return f"{feet}' {inches}'' / {height} cm"
        
        feet = int(height[0])
        inches = int(height[1:]) if len(height) > 1 else 0
        cm = int(feet * 30.48 + inches * 2.54)
        return f"{feet}' {inches}'' / {cm} cm"
    except ValueError:
        return "M"

def convert_weight(weight):
    """
    Converts weight to both pounds and kilograms.

    Args:
        weight (str): Client's weight in various formats.

    Returns:
        str: Weight in both lbs and kg.
    """
    weight = weight.replace(" ", "").strip()
    
    if weight in ["M", ""]:
        return "M"
    
    try:
        if "kg" in weight:
            weight_value = float(weight.replace('kg', ""))
            return f"{int(weight_value * 2.20462)} lbs / {weight_value} kg"
        elif "lbs" in weight:
            weight_value = float(weight.replace('lbs', ""))
            return f"{weight_value} lbs / {int(weight_value * 0.453592)} kg"
        
        weight_value = float(re.sub(r'[^0-9a-zA-Z]+', '', weight))
        return f"{weight_value} lbs / {int(weight_value * 0.453592)} kg"
    except ValueError:
        return "M"

def format_title_cells(title, data, title_cell):
    """
    Formats title cells by highlighting missing data fields.

    Args:
        title (str): The title of the field.
        data (Cell): The table cell containing data.
        title_cell (Cell): The table cell containing the title.
    """
    data_text = data.text.strip()
    
    if not data_text or data_text == "M":
        if not title_cell.text:
            title_cell.paragraphs[0].add_run(title).font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            title_cell.paragraphs[0].add_run().font.highlight_color = WD_COLOR_INDEX.YELLOW
        data.text = 'M'
    elif not title_cell.text:
        title_cell.paragraphs[0].add_run(title)


