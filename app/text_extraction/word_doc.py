from .constants import TITLES
from .utils import change_orientation
from .data_extraction import extract_customer_data
from .formatting_utils import convert_height, convert_weight, format_title_cells
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX
from pathlib import Path


class GuestDocumentGenerator:
    """
    A class to generate Word documents containing guest information tables.
    """
    
    def __init__(self, file_path):
        """
        Initializes the GuestDocumentGenerator with a specified file path.
        
        Args:
            file_path (str): The directory containing guest files.
        """
        self.file_path = file_path
        self.document = Document()
        self.set_page_format()
    
    def set_page_format(self):
        """
        Sets the page size to Letter (8.5" x 11") and ensures proper margins.
        """
        section = self.document.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
    
    def create_header(self, trip_name, start_date, end_date):
        """
        Creates a header for the document.
        
        Args:
            trip_name (str): The name of the trip.
            start_date (str): The start date of the trip.
            end_date (str): The end date of the trip.
        """
        p = self.document.add_paragraph('Guest Information Sheet')
        p.add_run(f'\nTrip Name: {trip_name}')
        p.add_run('\nIn Country Trip Dates: ')
        p.add_run(f'{start_date} thru {end_date}\t').font.highlight_color = WD_COLOR_INDEX.YELLOW
        p.add_run('\nCountries: MISSING').font.highlight_color = WD_COLOR_INDEX.YELLOW
    
    def create_table(self, page_num, num_rows=29, num_cols=4):
        """
        Creates a new table formatted to fit within a single page.
        
        Args:
            page_num (int): page number currently being processed
            num_rows (int): Number of rows in the table.
            num_cols (int): Number of columns in the table.
        
        Returns:
            Table: Empty table object to be populated with guest information
        """
        if page_num != 0:
            self.document.add_page_break()
        table = self.document.add_table(rows=num_rows, cols=num_cols)
        table.style = "Table Grid"
        return table

    def populate_table(self, table, guest_data, guest):
        """
        Populates table in word document with all the client information and makes
        a call to the format_title_cells function.

        Args:
            table (Table): table for storing client data
            guest_data (list): contains tuples of paired data (field, data)
            guest (int): the number of the guest (position)

        Returns:
            None
        """
        #Retrieve title cells:
        title_cells = table.columns[0].cells
        #Add client data:
        data_cells = table.columns[guest + 1].cells

        #Name:
        ind = 0
        data_cells[ind].text = guest_data.get('full name (as shown on passport):', 'M')
        ind += 1

        #Preferred Name:
        data_cells[ind].text = guest_data.get('preferred name:', 'M')
        ind += 1
        
        #Address:
        data_cells[ind].text = f"{str(guest_data.get('street address:', 'M'))} \
                \n{guest_data.get('city:', 'M')}, {guest_data.get('state/province:', 'M')}, {guest_data.get('zip code:', 'M')}\
                \n\n{guest_data.get('phone number:', 'M')}-Home"
        ind += 1
        
        #Birth date:
        response = guest_data.get('date of birth (mm/dd/yyyy):', 'M')
        data_cells[ind].text = response
        ind += 1
        
        #Age:
        data_cells[ind].text = guest_data.get('age at time of safari:', 'M')
        ind += 1
        
        #Gender:
        data_cells[ind].text = guest_data.get('gender', 'M')
        ind += 1

        #Email:
        data_cells[ind].text = guest_data.get('email:', 'M')
        ind += 1
        
        #Passport country/number:
        data_cells[ind].text = f"{guest_data.get('passport country:', 'M')} \n{guest_data.get('passport number:', 'M')}"
        ind += 1
        
        #Place of Issue:
        data_cells[ind].text = guest_data.get('passport place of issue:', 'M')
        ind += 1

        #Date of Issue:
        data_cells[ind].text = guest_data.get('date of issue:', 'M')
        ind += 1
        
        #Expiration Date:
        data_cells[ind].text = guest_data.get('passport expiration date:', 'M')
        ind += 1
        
        #Room Type:
        data_cells[ind].text = f"{guest_data.get('room type', 'M')} / {guest_data.get('bed type', 'M')}"
        ind += 1
        
        #Room Mate:
        data_cells[ind].text = guest_data.get('roommate', 'M')
        ind += 1
        
        #Emergency contact:
        data_cells[ind].text = f"{guest_data.get('full name:', 'M')} \n{guest_data.get('emergency email:', 'M')}" +\
                                f"\n{guest_data.get('phone number (include country code of outside usa):', 'M')}"
        ind += 1
        
        #Special Diet:
        data_cells[ind].text = guest_data.get('diet preference', 'M')
        ind += 1

        #Exclude from Diet:
        data_cells[ind].text = guest_data.get('diet exclusions', 'None')
        ind += 1
        
        #Allergies:
        data_cells[ind].text = f"Any allergies: {guest_data.get('any allergies', 'None')}" +\
                            f"\nAntibiotic allergies: {guest_data.get('antibiotic allergies', 'None')}"
        ind += 1

        #Allergies life threatening:
        data_cells[ind].text = f"Life threatening allergies: {guest_data.get('life threatening allergies', 'None')}"
        ind += 1
        
        #Medications:
        data_cells[ind].text = guest_data.get('medications', 'None')
        ind += 1
        
        #Blood type:
        data_cells[ind].text = f"{guest_data.get('blood type (if known):', 'M')}"
        ind += 1

        #Other Medical information:
        data_cells[ind].text = guest_data.get("equipment", 'M') + ' ' + guest_data.get("physical limitations", 'M')
        ind += 1

        #Fitness Level:
        data_cells[ind].text = guest_data.get('fitness', 'M')
        ind += 1

        #Travel/Medical insurance:
        data_cells[ind].text = f"{guest_data.get('Travel insurance:', 'M')}"
        ind += 1

        #Flying Doctors:
        data_cells[ind].text = f"{guest_data.get('Flying Doctors:', 'M')}"
        ind += 1

        #Special Occasions:
        data_cells[ind].text = guest_data.get('what is the occasion:', '') + " " + guest_data.get('celebrating', '')
        ind += 1

        #Height:
        data_cells[ind].text = convert_height(guest_data.get("height (ft' inch''):", 'M'))
        ind += 1
        
        #Weight:
        data_cells[ind].text = convert_weight(guest_data.get("weight (pounds):", '') + guest_data.get("weight (lbs):", ''))
        ind += 1
        
        #Shirt size:
        data_cells[ind].text = f"{guest_data.get('shirt size', 'M')}"
        ind += 1

        #Additional info:
        data_cells[ind].text = f"{guest_data.get('Additional info:', 'M')}"
        ind += 1

        ###------Add an extra row to the table if needed (could be based on a condition):-------###
        # row_cells = table.add_row().cells 
        # row_cells[0].text = 'Apple'
        # row_cells[1].text = '10'

        #Format all the cells based on input or lack of input:
        for title, data_cell, title_cell in zip(TITLES, data_cells, title_cells):
            format_title_cells(title, data_cell, title_cell)
        
        return None

    def generate_document(self, guest_files):
        """
        Generates a Word document containing guest information and saves it.
        
        Args:
            guest_files (list): List of guest file names.
        """
        # Rotate page if there are more than 2 guests:
        num_guests = len(guest_files)
        if num_guests > 2:
            change_orientation(self.document)

        # Extract customer file title and create header
        full_name = guest_files[0].split()[:2]
        trip_name = " ".join(full_name)
        self.create_header(trip_name, "MISSING", "MISSING")

        page_num = 0
        for guest_num, guest_file in enumerate(guest_files):
            if guest_num % 3 == 0:
                table = self.create_table(page_num)
                page_num += 1
            guest_info_dict = {}
            extract_customer_data(Path(f"{self.file_path}/{guest_file}"), guest_info_dict)
            self.populate_table(table, guest_info_dict, guest_num % 3)
    
        self.document.save(f'{self.file_path}/Guest Info {full_name[1]} {full_name[0]}.docx')



